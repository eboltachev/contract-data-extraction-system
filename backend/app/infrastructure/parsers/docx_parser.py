from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.domain.documents import DocumentFragment, ParsedDocument

KNOWN_SECTION_HEADINGS = (
    "предмет договора",
    "стоимость договора",
    "порядок оплаты и приемки работ",
    "сроки выполнения работ",
    "обеспечение строительства материалами",
    "права и обязанности подрядчика",
    "права и обязанности субподрядчика",
    "охранные мероприятия",
    "изменения в объеме работ",
    "организация производства работ",
    "гарантии качества",
    "форс-мажорные условия",
    "разрешение споров",
    "ответственность сторон",
    "конфиденциальность",
    "расторжение договора",
    "адреса и реквизиты сторон",
    "подписи сторон",
)


def parse_docx(path: Path) -> ParsedDocument:
    document = Document(path)
    fragments: list[DocumentFragment] = []
    tables: list[list[list[str]]] = []

    fragments.extend(_parse_headers_and_footers(document))
    body_fragments, body_tables = _parse_body(document)
    fragments.extend(body_fragments)
    tables.extend(body_tables)
    fragments = _deduplicate_fragments(fragments)

    return ParsedDocument(
        text="\n".join(fragment.text for fragment in fragments),
        tables=tables,
        fragments=fragments,
    )


def _parse_headers_and_footers(document: DocxDocument) -> list[DocumentFragment]:
    fragments: list[DocumentFragment] = []
    for section_idx, section in enumerate(document.sections, start=1):
        header_text = _join_paragraphs(section.header.paragraphs)
        if header_text:
            fragments.append(DocumentFragment(section=f"Колонтитул {section_idx}", clause=None, text=header_text))
        footer_text = _join_paragraphs(section.footer.paragraphs)
        if footer_text:
            fragments.append(DocumentFragment(section=f"Колонтитул {section_idx}", clause=None, text=footer_text))
    return fragments


def _parse_body(document: DocxDocument) -> tuple[list[DocumentFragment], list[list[list[str]]]]:
    fragments: list[DocumentFragment] = []
    tables: list[list[list[str]]] = []
    current_section = "Шапка договора"

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = _clean_text(block.text)
            if not text:
                continue

            if _is_section_heading(text, block.style.name if block.style else ""):
                current_section = _section_name(text)

            fragments.append(DocumentFragment(section=current_section, clause=_extract_clause(text), text=text))
            continue

        if isinstance(block, Table):
            table_rows: list[list[str]] = []
            for row in block.rows:
                cells = [_clean_text(cell.text) for cell in row.cells]
                if not any(cells):
                    continue
                table_rows.append(cells)
                line = " | ".join(cells)
                fragments.append(DocumentFragment(section=current_section, clause=None, text=line))

                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = _clean_text(paragraph.text)
                        if not text:
                            continue
                        fragments.append(DocumentFragment(section=current_section, clause=_extract_clause(text), text=text))
            if table_rows:
                tables.append(table_rows)

    return fragments, tables


def _iter_block_items(document: DocxDocument) -> Iterable[Paragraph | Table]:
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _join_paragraphs(paragraphs: Iterable[Paragraph]) -> str:
    return "\n".join(text for text in (_clean_text(paragraph.text) for paragraph in paragraphs) if text)


def _clean_text(value: str) -> str:
    value = value.replace("\xa0", " ").replace("\u200b", "")
    value = re.sub(r"[ \t\r\f\v]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _is_section_heading(text: str, style: str) -> bool:
    normalized = text.strip()
    normalized_key = _normalize_heading(normalized)
    style = (style or "").lower()

    if not normalized or "|" in normalized:
        return False
    if normalized_key in KNOWN_SECTION_HEADINGS:
        return True
    if any(normalized_key.startswith(heading) for heading in KNOWN_SECTION_HEADINGS):
        return True
    if "heading" in style or "заголов" in style:
        return True
    if len(normalized) > 180:
        return False
    if re.match(r"^\d{1,2}\.\s+[А-ЯЁA-Z].+", normalized):
        return True
    if re.match(r"^\d{1,2}\.\s*$", normalized):
        return True
    letters = [char for char in normalized if char.isalpha()]
    if 8 <= len(normalized) <= 160 and letters and all(char.isupper() for char in letters):
        return True
    return False


def _section_name(text: str) -> str:
    cleaned = text.strip().rstrip(".")
    key = _normalize_heading(cleaned)
    for heading in KNOWN_SECTION_HEADINGS:
        if key == heading or key.startswith(heading):
            return heading[:1].upper() + heading[1:]
    return cleaned[:180]


def _normalize_heading(value: str) -> str:
    value = value.lower().replace("ё", "е")
    value = re.sub(r"^\d{1,2}(?:\.\d+)*\.?\s+", "", value)
    value = re.sub(r"[^а-яa-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _extract_clause(text: str) -> str | None:
    match = re.match(r"^\s*(\d{1,2}(?:\.\d{1,3}){0,6})\.?\s+", text)
    return match.group(1) if match else None


def _deduplicate_fragments(fragments: list[DocumentFragment]) -> list[DocumentFragment]:
    seen: set[tuple[str | None, str | None, str]] = set()
    result: list[DocumentFragment] = []
    for fragment in fragments:
        key = (fragment.section, fragment.clause, fragment.text)
        if key in seen:
            continue
        seen.add(key)
        result.append(fragment)
    return result
